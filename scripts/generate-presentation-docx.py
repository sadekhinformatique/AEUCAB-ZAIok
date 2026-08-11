# -*- coding: utf-8 -*-
"""Génère la présentation Word de SGIAU (AEUCAB-ZAI).

Usage :  .venv/Scripts/python scripts/generate-presentation-docx.py [nom_sortie]
Sortie : Presentation-SGIAU.docx à la racine du workspace (ou le nom fourni).
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------- palette
# Couleurs du logo de l'amicale : vert #086808, rouge #B80808, noir #101010
VERT      = RGBColor(0x08, 0x68, 0x08)   # vert du logo
VERT_HEX  = "086808"
ROUGE     = RGBColor(0xB8, 0x08, 0x08)   # rouge du logo
VERT_CLAIR = "E5F3E5"                    # fond vert clair (tableaux)
GRIS      = RGBColor(0x55, 0x5F, 0x6B)
NOIR      = RGBColor(0x10, 0x10, 0x10)
BLANC     = RGBColor(0xFF, 0xFF, 0xFF)
LOGO_PATH = "AEUCAB-ZAIok/public/logo-aeucab.png"

ROLE_STYLE = [
    # (identifiant, libellé, rgb, hex_fond)
    ("admin", "ADMIN_IT — Administrateur informatique", RGBColor(0x47, 0x55, 0x69), "E2E8F0"),
    ("president", "PRESIDENT — Président", RGBColor(0xB4, 0x53, 0x09), "FEF3C7"),
    ("secretaire", "SECRETAIRE — Secrétaire général", RGBColor(0x04, 0x78, 0x57), "D1FAE5"),
    ("tresorier", "TRESORIER — Trésorier", RGBColor(0xBE, 0x12, 0x3C), "FFE4E6"),
    ("caissier", "CAISSIER — Caissier", RGBColor(0x0E, 0x74, 0x90), "CFFAFE"),
    ("commissaire", "COMMISSAIRE — Commissaire aux comptes", RGBColor(0x6D, 0x28, 0xD9), "EDE9FE"),
    ("member", "MEMBER — Membre", RGBColor(0x52, 0x64, 0x33), "ECFCCB"),
    ("custom", "CUSTOM — Rôle personnalisé", RGBColor(0x86, 0x1E, 0x8C), "FAE8FF"),
]

ROLES_DETAIL = {
    "admin": (
        "Rôle technique : maîtrise totale du système.",
        ["Gère les utilisateurs : création, rôles, activation, réinitialisation des mots de passe.",
         "Administre la sécurité : verrouillages, sessions, politique de mot de passe fort.",
         "Consulte le journal d'audit et configure la synchronisation et l'import/export.",
         "Crée les comptes du bureau exécutif et applique la politique de sécurité."],
    ),
    "president": (
        "Rôle décisionnel : pilote l'amicale et a le dernier mot sur les validations.",
        ["Valide en dernier ressort les adhésions (après le secrétaire) : étape « Validation président ».",
         "Valide les dépenses et les opérations financières importantes.",
         "Suit le tableau de bord, les statistiques et la trésorerie globale.",
         "Peut consulter, ajouter, modifier, supprimer, valider, exporter et imprimer."],
    ),
    "secretaire": (
        "Rôle administratif : garant de la vie documentaire et des procédures.",
        ["Première validation des adhésions (étape « Validé par le secrétaire »).",
         "Tient le registre des membres ; rédige les PV de réunions et les ordres du jour.",
         "Gère les documents, archives, annonces et notifications.",
         "Organise les présences (manuel + QR code). Crée, modifie et valide — ne supprime pas."],
    ),
    "tresorier": (
        "Rôle financier : responsable de la comptabilité et du suivi budgétaire.",
        ["Gère les cotisations : types, montants, encaissements.",
         "Tient la comptabilité : journal, plan comptable, balance, exercices, clôture.",
         "Crée et valide les dépenses ; suit les reçus et les mouvements de caisse.",
         "Crée, modifie et valide — ne supprime pas (traçabilité financière)."],
    ),
    "caissier": (
        "Rôle opérationnel : encaisse les paiements au quotidien.",
        ["Encaissement des cotisations : espèces, mobile money, virement bancaire.",
         "Émet et imprime les reçus (PDF + QR code de vérification).",
         "Gère la caisse : entrées, sorties, transferts entre comptes.",
         "Consulte et ajoute, imprime les reçus — ne modifie pas, ne valide pas, ne supprime pas "
         "(toute correction passe par le trésorier)."],
    ),
    "commissaire": (
        "Rôle de contrôle : vérifie la régularité des opérations (contre-pouvoir).",
        ["Consulte l'ensemble des données : membres, finances, caisse, dépenses, votes.",
         "Exporte et imprime pour produire des rapports de contrôle.",
         "Lecture seule garantie : ne crée, ne modifie, ne valide et ne supprime rien."],
    ),
    "member": (
        "Rôle adhérent : accès limité à l'espace membre (application mobile web).",
        ["Consulte les annonces, sa cotisation et ses reçus.",
         "Peut soumettre des demandes.",
         "Aucun accès au back-office."],
    ),
    "custom": (
        "Rôle vierge destiné à être paramétré selon les besoins spécifiques de l'amicale.",
        ["Lecture seule par défaut.",
         "Exemples d'usage : responsable communication, responsable sport, etc."],
    ),
}

PERMS = ["Consulter", "Ajouter", "Modifier", "Supprimer", "Valider", "Exporter", "Imprimer"]
PERM_MATRIX = {
    "admin":       [True, True, True, True, True, True, True],
    "president":   [True, True, True, True, True, True, True],
    "secretaire":  [True, True, True, False, True, True, True],
    "tresorier":   [True, True, True, False, True, True, True],
    "caissier":    [True, True, False, False, False, False, True],
    "commissaire": [True, False, False, False, False, True, True],
    "member":      [True, False, False, False, False, False, False],
    "custom":      [True, False, False, False, False, False, False],
}

MODULES = [
    ("Pilotage", ["Tableau de bord", "Statistiques", "Recherche globale"]),
    ("Membres", ["Membres", "Adhésions", "Cartes membres", "Espace membre"]),
    ("Finances", ["Cotisations", "Reçus", "Comptabilité", "Caisse", "Dépenses"]),
    ("Vie associative", ["Activités", "Réunions", "Présences", "Élections", "Votes internes"]),
    ("Ressources", ["Documents", "Inventaire", "Formations", "Bibliothèque", "Partenaires", "Archives"]),
    ("Système", ["Notifications", "Import / Export", "Utilisateurs & sécurité", "Journal d'audit", "Synchronisation"]),
]

COMPTES = [
    ("admin", "ADMIN_IT", "admin@sgiau.local"),
    ("president", "PRESIDENT", "president@ucab.com"),
    ("secretaire", "SECRETAIRE", "secretaire@ucab.com"),
    ("tresorier", "TRESORIER", "tresorier@ucab.com"),
    ("caissier", "CAISSIER", "caissier@ucab.com"),
    ("commissaire", "COMMISSAIRE", "commissaire@ucab.com"),
]

# ---------------------------------------------------------------- helpers
def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, color=None, size=10, align=None, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if white:
        r.font.color.rgb = BLANC
    elif color is not None:
        r.font.color.rgb = color

def add_page_number(paragraph):
    run = paragraph.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(it); run._r.append(f2)
    run.font.size = Pt(9)
    run.font.color.rgb = GRIS

def h1(doc, text, color=VERT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = color
    return p

def h2(doc, text, color=VERT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = color
    return p

def para(doc, text, size=11, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return p

def bullets(doc, items, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(it)
        r.font.size = Pt(size)

# ---------------------------------------------------------------- document
doc = Document()

# Marges
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Style par défaut
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = NOIR

# Pied de page : nom + numéro
footer_p = doc.sections[0].footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run("SGIAU — Système de Gestion Intégrée de l'Amicale Universitaire   ·   page ")
fr.font.size = Pt(9); fr.font.color.rgb = GRIS
add_page_number(footer_p)

# ================================================================ PAGE DE GARDE
for _ in range(3):
    doc.add_paragraph()

# Logo de l'amicale (vert / rouge / noir)
if os.path.exists(LOGO_PATH):
    doc.add_picture(LOGO_PATH, width=Cm(3.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[ Logo de l'amicale ]"); r.font.size = Pt(11); r.font.color.rgb = GRIS

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Université Cheikh Ahmadou Bamba (UCAB)")
r.bold = True; r.font.size = Pt(19); r.font.color.rgb = VERT

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
r = p.add_run("Amicale des Étudiants — AEUCAB-ZAI")
r.font.size = Pt(12); r.font.color.rgb = GRIS

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
r = p.add_run("SGIAU"); r.bold = True; r.font.size = Pt(52); r.font.color.rgb = VERT

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Système de Gestion Intégrée de l'Amicale Universitaire")
r.font.size = Pt(17); r.font.color.rgb = GRIS

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
r = p.add_run("Présentation du système et des rôles")
r.font.size = Pt(13); r.font.color.rgb = GRIS

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(26)
r = p.add_run("contact@ucab.sn   ·   +(221) 776458011   ·   www.ucab.sn   ·   www.aeucab-zai.vercel.app")
r.font.size = Pt(10); r.font.color.rgb = GRIS

doc.add_page_break()

# ================================================================ 1. PRÉSENTATION
h1(doc, "1.  Présentation générale")
para(doc, "SGIAU est une application web complète qui centralise la gestion de l'amicale "
          "universitaire : membres, adhésions, cotisations, comptabilité, caisse, dépenses, "
          "activités, réunions, élections, documents, inventaire, formations, bibliothèque, "
          "partenaires, archives, notifications et sécurité.")
para(doc, "L'application offre deux espaces :", bold=True, space_after=2)
bullets(doc, [
    "Back-office : gestion complète par le bureau exécutif et l'administrateur informatique.",
    "Espace membre : application mobile web où chaque membre consulte annonces, cotisations, reçus et demandes.",
])
para(doc, "L'application est déployée en production : https://aeucab-zai.vercel.app "
          "(vérification d'état disponible sur /api/health).")

h2(doc, "Les 28 modules, organisés en 6 familles")
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
widths = (Cm(5.2), Cm(11.6))
hdr = t.rows[0].cells
for i, txt in enumerate(("Famille", "Modules")):
    cell_text(hdr[i], txt, bold=True, white=True, size=11)
    shade_cell(hdr[i], VERT_HEX)
for fam, mods in MODULES:
    row = t.add_row().cells
    cell_text(row[0], fam, bold=True, size=10.5)
    cell_text(row[1], " · ".join(mods), size=10.5)
    shade_cell(row[0], VERT_CLAIR)
for row in t.rows:
    for i, c in enumerate(row.cells):
        c.width = widths[i]

# ================================================================ 2. L'UNIVERSITÉ
h1(doc, "2.  L'Université Cheikh Ahmadou Bamba (UCAB)")
para(doc, "L'UCAB est un établissement privé d'enseignement supérieur doté d'une mission de "
          "service public d'éducation, de formation et de recherche. Elle s'appuie sur les valeurs "
          "islamiques du travail rédempteur, de la paix, de la tolérance et de la quête du savoir, "
          "conformément à l'enseignement de Cheikh Ahmadou Bamba, et exerce sa mission en toute "
          "conformité avec les lois et règlements de la République du Sénégal.")
para(doc, "Président et Responsable Moral : Serigne Mame Mor MBACKE Mourtada  ·  "
          "devise « L'Excellence au service du savoir »", bold=True)

h2(doc, "2.1  Les 4 centres et leurs facultés")
centres = [
    ("Centre Cheikh Mouhamadou Mourtada Mbacké — Dakar",
     "Faculté des Sciences Islamiques et Technologies",
     "Dakar Mbao, face LGI Mbao · dakarcentre@ucab.sn · dakar.ucab.sn"),
    ("Centre Serigne Mame Mor Diarra Mbacké — Saint-Louis",
     "Faculté des Technologies Agro-Alimentaires, des Sciences Économiques et Sociales",
     "Saint-Louis · stlouiscentre@ucab.sn · stlouisucab.sn"),
    ("Centre de Touba Darou Alim (siège de l'université)",
     "Faculté des Sciences Religieuses, des Humanités et Civilisations",
     "Touba Darou Alim · toubacentre@ucab.sn · touba.ucab.sn"),
    ("Centre Serigne Saliou Mbacké — Bambey",
     "Faculté de Développement Rural",
     "Bambey · bambeycentre@ucab.sn · bambey.ucab.sn"),
]
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
widths = (Cm(6.2), Cm(5.4), Cm(5.2))
hdr = t.rows[0].cells
for i, txt in enumerate(("Centre", "Faculté", "Localisation & contact")):
    cell_text(hdr[i], txt, bold=True, white=True, size=10)
    shade_cell(hdr[i], VERT_HEX)
for centre, fac, contact in centres:
    row = t.add_row().cells
    cell_text(row[0], centre, bold=True, size=9.5)
    cell_text(row[1], fac, size=9.5)
    cell_text(row[2], contact, size=9)
    shade_cell(row[0], VERT_CLAIR)
for row in t.rows:
    for i, c in enumerate(row.cells):
        c.width = widths[i]

h2(doc, "2.2  Les formations (licences)")
bullets(doc, [
    "Sciences Religieuses ;  Histoire et Civilisations ;",
    "Électromécanique ;  Électricité et Informatique Industrielle ;  Informatique de Gestion ;",
    "Technologies Agroalimentaires ;  Développement Rural (Agriculture ou Élevage) ;",
    "Administration (Finances Islamiques, Comptabilité, GRH, Marketing & Communication) ;",
    "En projet : Journalisme & Communication, Tourisme Religieux, Géomètre-Topographie, Génie Civil, "
    "Masters en Diplomatie & Relations Internationales et en Sciences Islamiques.",
])

h2(doc, "2.3  Organisation administrative")
bullets(doc, [
    "Instances de gouvernance : Conseil d'administration, Conseil académique, Conseil d'établissement, "
    "Conseil pédagogique, Conseil de département, Conseil de discipline.",
    "Organes : Président, Vice-Président, Agence comptable, Secrétariat général, services "
    "(Communication, RH, Finances, Coopération), directions (Scolarité & Système d'Information, "
    "Études & Pédagogie), puis les facultés et leurs départements.",
])

# ================================================================ 3. ARCHITECTURE
h1(doc, "3.  Architecture technique")
arch = [
    ("Framework", "Next.js 16 (App Router)"),
    ("Langage", "TypeScript 5"),
    ("Interface", "Tailwind CSS 4 + shadcn/ui + icônes Lucide"),
    ("Graphiques", "Recharts"),
    ("ORM", "Prisma 6"),
    ("Base de données", "Neon PostgreSQL (cloud)"),
    ("État applicatif", "Zustand + TanStack Query"),
    ("Authentification", "Sessions signées HMAC-SHA256, mots de passe hachés bcrypt"),
]
t = doc.add_table(rows=0, cols=2)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
widths = (Cm(5.2), Cm(11.6))
for k, v in arch:
    row = t.add_row().cells
    cell_text(row[0], k, bold=True, size=10.5)
    cell_text(row[1], v, size=10.5)
    shade_cell(row[0], VERT_CLAIR)
for row in t.rows:
    for i, c in enumerate(row.cells):
        c.width = widths[i]

# ================================================================ 4. GUIDE D'UTILISATION
h1(doc, "4.  Guide d'utilisation")

h2(doc, "4.1  Connexion")
bullets(doc, [
    "Ouvrir https://aeucab-zai.vercel.app (redirection automatique vers la page de connexion).",
    "Saisir l'identifiant (ex. president) et le mot de passe.",
    "Premier login avec un mot de passe temporaire : l'application impose la création d'un "
    "mot de passe personnel (8 caractères minimum, majuscule, minuscule, chiffre et caractère spécial).",
    "Sécurité : après 5 tentatives échouées, le compte est verrouillé 15 minutes.",
])

h2(doc, "4.2  Navigation")
bullets(doc, [
    "Menu latéral : les 28 modules sont regroupés en 6 familles (Pilotage, Membres, Finances, "
    "Vie associative, Ressources, Système).",
    "Tableau de bord : vue d'ensemble en temps réel (membres, finances, alertes, graphiques).",
    "Recherche globale : recherche multicritère sur toutes les entités.",
    "Le rôle du compte détermine les modules accessibles et les actions possibles.",
])

h2(doc, "4.3  Parcours types")
for titre, etapes in [
    ("Un nouveau membre adhère", [
        "Le dossier apparaît « En attente » dans le module Adhésions.",
        "Le secrétaire valide (étape « Validé par le secrétaire »).",
        "Le président approuve (étape « Validation président »).",
        "Le membre devient actif ; une carte membre avec QR code est générée.",
    ]),
    ("Le membre paie sa cotisation", [
        "Module Cotisations : choisir le type (annuelle, mensuelle, don…).",
        "Encaisser le paiement (espèces, mobile money, banque).",
        "Un reçu PDF avec QR code est généré automatiquement.",
        "La comptabilité et la caisse sont mises à jour en temps réel.",
    ]),
    ("Une dépense est engagée", [
        "Création de la dépense par le responsable (module Dépenses).",
        "Validation par le trésorier ou le président.",
        "Écritures comptables automatiques et mouvement de caisse.",
    ]),
    ("Une élection / un vote", [
        "Module Élections : déclaration des candidats par poste.",
        "Vote électronique — un vote par membre, QR code de vérification.",
        "Résultats publiés ; votes internes anonymes disponibles également.",
    ]),
]:
    para(doc, "▶ " + titre, bold=True, space_after=2)
    bullets(doc, etapes)

doc.add_page_break()

# ================================================================ 5. RÔLES
h1(doc, "5.  Les rôles et leur travail dans le système")
para(doc, "Huit rôles sont définis. Chaque rôle détermine les modules accessibles et les "
          "actions possibles : consulter, ajouter, modifier, supprimer, valider, exporter, imprimer.")

h2(doc, "5.1  Matrice des permissions")
t = doc.add_table(rows=1, cols=1 + len(PERMS))
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
widths = [Cm(3.9)] + [Cm(1.85)] * len(PERMS)
hdr = t.rows[0].cells
cell_text(hdr[0], "Rôle", bold=True, white=True, size=9.5)
shade_cell(hdr[0], VERT_HEX)
for i, perm in enumerate(PERMS, start=1):
    cell_text(hdr[i], perm, bold=True, white=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(hdr[i], VERT_HEX)
for key, label, _, hexb in ROLE_STYLE:
    row = t.add_row().cells
    cell_text(row[0], label, bold=True, size=9)
    shade_cell(row[0], hexb)
    for i, val in enumerate(PERM_MATRIX[key], start=1):
        cell_text(row[i], "✓" if val else "—", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=VERT if val else GRIS)
for row in t.rows:
    for i, c in enumerate(row.cells):
        c.width = widths[i]

para(doc, "")
para(doc, "Remarque : la matrice ci-dessus est celle du module « Utilisateurs & sécurité ». "
          "Le contrôle strict d'administrateur (requireAdmin) protège actuellement les routes "
          "utilisateurs et journal d'audit ; les autres modules reposent sur une session valide. "
          "Une montée en rigueur RBAC est recommandée (voir AUDIT_SGIAU.md).", size=9, color=GRIS)

h2(doc, "5.2  Fiches détaillées par rôle")
for key, label, rgb, hexb in ROLE_STYLE:
    # Encadré titre du rôle
    tt = doc.add_table(rows=1, cols=1)
    tt.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tt.rows[0].cells[0]
    c.width = Cm(16.8)
    cell_text(c, label, bold=True, white=True, size=12)
    shade_cell(c, "%02X%02X%02X" % (rgb[0], rgb[1], rgb[2]))
    intro, missions = ROLES_DETAIL[key]
    para(doc, intro, bold=True, space_after=2)
    bullets(doc, missions)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# ================================================================ 6. SÉCURITÉ
h1(doc, "6.  Sécurité")
bullets(doc, [
    "Mots de passe hachés en bcrypt (coût 12) — jamais stockés en clair.",
    "Sessions sécurisées : cookies httpOnly / SameSite / Secure, jetons HMAC-SHA256 à durée limitée.",
    "Verrouillage automatique après 5 échecs de connexion (15 minutes).",
    "Politique de mot de passe fort appliquée à la création et au changement.",
    "Changement de mot de passe forcé au premier login (mot de passe temporaire).",
    "Journal d'audit : chaque action (création, modification, suppression, validation, "
    "export, impression) est tracée avec état avant / après en JSON.",
    "Seed de démonstration désactivé en production (HTTP 403).",
    "En-têtes de sécurité HTTP (nosniff, X-Frame-Options, Referrer-Policy, Permissions-Policy).",
])

# ================================================================ 7. ACCÈS & COMPTES
h1(doc, "7.  Accès et comptes")
h2(doc, "7.1  Adresses")
acc = [
    ("Application web (production)", "https://aeucab-zai.vercel.app"),
    ("Développement local", "http://localhost:3000"),
    ("Espace membre (mobile web)", "https://aeucab-zai.vercel.app/member-space"),
    ("Vérification d'état", "https://aeucab-zai.vercel.app/api/health"),
]
t = doc.add_table(rows=0, cols=2)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
widths = (Cm(7.4), Cm(9.4))
for k, v in acc:
    row = t.add_row().cells
    cell_text(row[0], k, bold=True, size=10.5)
    cell_text(row[1], v, size=10.5)
    shade_cell(row[0], VERT_CLAIR)
for row in t.rows:
    for i, c in enumerate(row.cells):
        c.width = widths[i]

h2(doc, "7.2  Comptes en base (production)")
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
widths = (Cm(4.0), Cm(5.6), Cm(7.2))
hdr = t.rows[0].cells
for i, txt in enumerate(("Identifiant", "Rôle", "Email")):
    cell_text(hdr[i], txt, bold=True, white=True, size=10.5)
    shade_cell(hdr[i], VERT_HEX)
for u, role, email in COMPTES:
    row = t.add_row().cells
    cell_text(row[0], u, bold=True, size=10)
    cell_text(row[1], role, size=10)
    cell_text(row[2], email, size=10)
for row in t.rows:
    for i, c in enumerate(row.cells):
        c.width = widths[i]

para(doc, "")
para(doc, "Les mots de passe ont été réinitialisés (mots de passe temporaires affichés une seule "
          "fois, changement forcé à la première connexion). Ils ne figurent pas dans ce document. "
          "En cas de perte, un administrateur peut réinitialiser un mot de passe depuis le module "
          "« Utilisateurs & sécurité ».", size=9, color=GRIS)

OUT = sys.argv[1] if len(sys.argv) > 1 else "Presentation-SGIAU.docx"
doc.save(OUT)
print(f"OK — {OUT} généré")

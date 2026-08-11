# -*- coding: utf-8 -*-
"""Génère le flyer d'inscription Word (AEUCAB-ZAI) avec logo et couleurs.

Usage :  .venv/Scripts/python scripts/generate-flyer-docx.py [nom_sortie]
Sortie : Flyer-Inscription-SGIAU.docx à la racine du workspace (ou le nom fourni).
"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------- palette
# Couleurs du logo de l'amicale : vert #086808, rouge #B80808, noir #101010
VERT      = RGBColor(0x08, 0x68, 0x08)
VERT_HEX  = "086808"
VERT_CLAIR = "E5F3E5"
VERT_FONCE = "053F05"
ROUGE     = RGBColor(0xB8, 0x08, 0x08)
ROUGE_HEX = "B80808"
NOIR      = RGBColor(0x10, 0x10, 0x10)
GRIS      = RGBColor(0x55, 0x5F, 0x6B)
BLANC     = RGBColor(0xFF, 0xFF, 0xFF)
LOGO_PATH = "AEUCAB-ZAIok/public/logo-aeucab.png"

URL_APP = "https://aeucab-zai.vercel.app/espace-membre"

# ---------------------------------------------------------------- helpers
def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, color=None, size=10, align=None, white=False, space_after=1):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    for i, chunk in enumerate(text.split("\n")):
        if i > 0:
            p = cell.add_paragraph()
            p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(chunk)
        r.bold = bold
        r.font.size = Pt(size)
        if white:
            r.font.color.rgb = BLANC
        elif color is not None:
            r.font.color.rgb = color

def full_width_table(doc, cols=1):
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    return t

def set_widths(table, widths):
    for row in table.rows:
        for i, c in enumerate(row.cells):
            c.width = widths[i]

# ================================================================ document
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = NOIR

# ------------------------------------------------------------- bandeau haut (vert)
band = full_width_table(doc)
row = band.add_row().cells
shade_cell(row[0], VERT_HEX)
para = row[0].paragraphs[0]
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.add_run()
run.add_picture(LOGO_PATH, height=Cm(2.1))
p2 = row[0].add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
r2 = p2.add_run("UNIVERSITÉ CHEIKH AHMADOU BAMBA (UCAB)")
r2.bold = True; r2.font.size = Pt(13); r2.font.color.rgb = BLANC
p3 = row[0].add_paragraph()
r3 = p3.add_run("Amicale des Étudiants — AEUCAB-ZAI")
r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(0xD9, 0xF2, 0xD9)

# ------------------------------------------------------------- liseré rouge
band2 = full_width_table(doc)
row2 = band2.add_row().cells
shade_cell(row2[0], ROUGE_HEX)
row2[0].paragraphs[0].text = ""

# ------------------------------------------------------------- titre
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("🎓 Rejoignez l'amicale !")
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = VERT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Inscrivez-vous à l'espace membre — 2 minutes depuis votre téléphone")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = ROUGE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
r = p.add_run("Fini les files d'attente et les formulaires papier : votre adhésion, "
              "votre carte et vos reçus, directement sur votre mobile.")
r.font.size = Pt(11); r.font.color.rgb = GRIS

# ------------------------------------------------------------- avantages
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("✅ Ce que vous gagnez en vous inscrivant")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = VERT

avantages = [
    "Adhésion simplifiée — votre demande part automatiquement au bureau, et vous suivez sa validation en temps réel",
    "Carte de membre numérique avec QR code, toujours avec vous sur votre téléphone",
    "Cotisations en toute simplicité — payez et recevez votre reçu officiel sans passer par le secrétariat",
    "Restez informé — annonces, activités, réunions et événements en un coup d'œil",
    "Vos documents en ligne — attestations, certificats et demandes sans déplacement",
    "Un espace 100 % personnel et sécurisé, accessible à tout moment",
]
t = full_width_table(doc, cols=1)
for av in avantages:
    row = t.add_row().cells
    cell_text(row[0], "✔  " + av, size=11, color=NOIR)
    shade_cell(row[0], VERT_CLAIR)

# ------------------------------------------------------------- étapes
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("📲 Comment s'inscrire ?")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = VERT

etapes = [
    ("1", "Ouvrez l'application « Espace membre »"),
    ("2", "Touchez « Pas encore membre ? S'inscrire »"),
    ("3", "Remplissez vos informations (2 minutes) et choisissez votre identifiant et mot de passe"),
    ("4", "C'est fait ! Votre adhésion est enregistrée — le bureau la valide rapidement"),
]
t = full_width_table(doc, cols=2)
for num, txt in etapes:
    row = t.add_row().cells
    cell_text(row[0], num, bold=True, white=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(row[0], VERT_HEX)
    cell_text(row[1], txt, size=11)
set_widths(t, (Cm(1.4), Cm(15.4)))

# ------------------------------------------------------------- encadré CTA
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(2)
cta = full_width_table(doc)
crow = cta.add_row().cells
shade_cell(crow[0], VERT_FONCE)
cell_text(crow[0], "Rejoignez la grande famille de l'amicale — votre engagement commence ici ! 💚",
          bold=True, white=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
p = crow[0].add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(URL_APP)
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = RGBColor(0xFF, 0xE0, 0x8A)
p = crow[0].add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Gratuit · Rapide · Sécurisé")
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xD9, 0xF2, 0xD9)

# ------------------------------------------------------------- pied de page
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(2)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Mot de passe oublié ? La fonction « Mot de passe oublié » vous permet de le récupérer en un instant.")
r.font.size = Pt(10); r.font.color.rgb = GRIS

foot = full_width_table(doc)
frow = foot.add_row().cells
shade_cell(frow[0], ROUGE_HEX)
cell_text(frow[0],
          "Université Cheikh Ahmadou Bamba (UCAB) — 4 centres : Dakar · Saint-Louis · Touba · Bambey\n"
          "contact@ucab.sn  ·  +(221) 776458011  ·  www.ucab.sn",
          white=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

OUT = sys.argv[1] if len(sys.argv) > 1 else "Flyer-Inscription-SGIAU.docx"
doc.save(OUT)
print(f"OK — {OUT} généré")
